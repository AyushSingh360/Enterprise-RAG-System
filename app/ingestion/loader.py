"""
Document loaders for PDF, DOCX, and Markdown files.

Features:
- Extracts text with metadata preservation
- Tracks page numbers, sections, and source info
- Returns standardized Document objects
"""

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Any
from datetime import datetime, timezone
import structlog

logger = structlog.get_logger(__name__)


@dataclass
class Document:
    """
    Represents a loaded document with text and metadata.
    
    Metadata is preserved throughout the pipeline for citation.
    """
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    
    @property
    def source(self) -> str:
        """Get document source."""
        return self.metadata.get("source", "unknown")
    
    @property
    def page_number(self) -> Optional[int]:
        """Get page number if available."""
        return self.metadata.get("page_number")
    
    @property
    def section(self) -> Optional[str]:
        """Get section name if available."""
        return self.metadata.get("section")


class BaseLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, source: str | Path, **kwargs) -> list[Document]:
        """
        Load documents from a source.
        
        Args:
            source: File path or content string.
            **kwargs: Additional loader-specific arguments.
            
        Returns:
            List of Document objects with preserved metadata.
        """
        pass
    
    def _create_base_metadata(
        self,
        source: str,
        doc_type: str,
        **extra: Any,
    ) -> dict[str, Any]:
        """Create base metadata dict for a document."""
        return {
            "source": source,
            "document_type": doc_type,
            "loaded_at": datetime.now(timezone.utc).isoformat(),
            **extra,
        }


class PDFLoader(BaseLoader):
    """
    Loader for PDF documents.
    
    Extracts text page by page with page number metadata.
    """
    
    def load(self, source: str | Path, **kwargs) -> list[Document]:
        """
        Load a PDF file.
        
        Args:
            source: Path to PDF file.
            
        Returns:
            List of Documents, one per page with page_number metadata.
        """
        from pypdf import PdfReader
        
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {source}")
        
        documents: list[Document] = []
        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            
            if not text or not text.strip():
                logger.debug("Skipping empty page", page=page_num, file=path.name)
                continue
            
            # Clean text
            text = self._clean_text(text)
            
            metadata = self._create_base_metadata(
                source=path.name,
                doc_type="pdf",
                file_path=str(path.absolute()),
                page_number=page_num,
                total_pages=total_pages,
            )
            
            documents.append(Document(text=text, metadata=metadata))
        
        logger.info(
            "Loaded PDF",
            file=path.name,
            total_pages=total_pages,
            extracted_pages=len(documents),
        )
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted PDF text."""
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


class DOCXLoader(BaseLoader):
    """
    Loader for DOCX documents.
    
    Extracts text with section detection based on heading styles.
    """
    
    def load(self, source: str | Path, **kwargs) -> list[Document]:
        """
        Load a DOCX file.
        
        Args:
            source: Path to DOCX file.
            
        Returns:
            List of Documents with section metadata.
        """
        from docx import Document as DocxDocument
        from docx.opc.exceptions import PackageNotFoundError
        
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {source}")
        
        try:
            doc = DocxDocument(str(path))
        except PackageNotFoundError:
            raise ValueError(f"Invalid DOCX file: {source}")
        
        documents: list[Document] = []
        current_section: Optional[str] = None
        section_texts: list[str] = []
        section_start_para = 0
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect headings for section tracking
            is_heading = para.style.name.startswith("Heading")
            
            if is_heading:
                # Save previous section if exists
                if section_texts:
                    section_text = "\n\n".join(section_texts)
                    metadata = self._create_base_metadata(
                        source=path.name,
                        doc_type="docx",
                        file_path=str(path.absolute()),
                        section=current_section,
                        paragraph_start=section_start_para,
                        paragraph_end=i - 1,
                    )
                    documents.append(Document(text=section_text, metadata=metadata))
                
                # Start new section
                current_section = text
                section_texts = []
                section_start_para = i
            else:
                section_texts.append(text)
        
        # Save final section
        if section_texts:
            section_text = "\n\n".join(section_texts)
            metadata = self._create_base_metadata(
                source=path.name,
                doc_type="docx",
                file_path=str(path.absolute()),
                section=current_section,
            )
            documents.append(Document(text=section_text, metadata=metadata))
        
        # If no sections found, return as single document
        if not documents:
            all_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if all_text:
                metadata = self._create_base_metadata(
                    source=path.name,
                    doc_type="docx",
                    file_path=str(path.absolute()),
                )
                documents.append(Document(text=all_text, metadata=metadata))
        
        logger.info(
            "Loaded DOCX",
            file=path.name,
            sections=len(documents),
        )
        
        return documents


class MarkdownLoader(BaseLoader):
    """
    Loader for Markdown content.
    
    Extracts text with section detection based on headers.
    Converts markdown to plain text while preserving structure.
    """
    
    # Regex for markdown headers
    HEADER_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
    
    def load(
        self,
        source: str | Path,
        is_file: bool = True,
        source_name: Optional[str] = None,
        **kwargs,
    ) -> list[Document]:
        """
        Load markdown content.
        
        Args:
            source: Path to markdown file or raw markdown content.
            is_file: If True, source is a file path. If False, source is content.
            source_name: Optional name for the source when content is provided.
            
        Returns:
            List of Documents with section metadata.
        """
        if is_file:
            path = Path(source)
            if not path.exists():
                raise FileNotFoundError(f"Markdown file not found: {source}")
            
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            
            file_name = path.name
        else:
            content = str(source)
            file_name = source_name or "markdown_content"
        
        # Parse sections based on headers
        documents = self._parse_sections(content, file_name)
        
        if not documents:
            # No headers found, return as single document
            plain_text = self._markdown_to_text(content)
            if plain_text.strip():
                metadata = self._create_base_metadata(
                    source=file_name,
                    doc_type="markdown",
                )
                documents.append(Document(text=plain_text, metadata=metadata))
        
        logger.info(
            "Loaded Markdown",
            source=file_name,
            sections=len(documents),
        )
        
        return documents
    
    def _parse_sections(self, content: str, source_name: str) -> list[Document]:
        """Parse markdown into sections based on headers."""
        documents: list[Document] = []
        
        # Find all headers
        headers = list(self.HEADER_PATTERN.finditer(content))
        
        if not headers:
            return documents
        
        for i, match in enumerate(headers):
            header_level = len(match.group(1))
            header_text = match.group(2).strip()
            
            # Get section content
            start = match.end()
            end = headers[i + 1].start() if i + 1 < len(headers) else len(content)
            
            section_content = content[start:end].strip()
            
            if section_content:
                plain_text = self._markdown_to_text(section_content)
                
                if plain_text.strip():
                    metadata = self._create_base_metadata(
                        source=source_name,
                        doc_type="markdown",
                        section=header_text,
                        section_level=header_level,
                    )
                    documents.append(Document(text=plain_text, metadata=metadata))
        
        return documents
    
    def _markdown_to_text(self, markdown: str) -> str:
        """Convert markdown to plain text."""
        import markdown as md
        from bs4 import BeautifulSoup
        
        # Convert to HTML then extract text
        html = md.markdown(markdown, extensions=['tables', 'fenced_code'])
        soup = BeautifulSoup(html, "html.parser")
        
        # Get text with spacing
        text = soup.get_text(separator="\n")
        
        # Clean up
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()


class DocumentLoader:
    """
    Factory for document loaders.
    
    Provides unified interface for loading different document types.
    """
    
    def __init__(self):
        """Initialize with available loaders."""
        self._loaders: dict[str, BaseLoader] = {
            "pdf": PDFLoader(),
            "docx": DOCXLoader(),
            "markdown": MarkdownLoader(),
            "md": MarkdownLoader(),
        }
    
    def load(
        self,
        source: str | Path,
        doc_type: Optional[str] = None,
        **kwargs,
    ) -> list[Document]:
        """
        Load a document using the appropriate loader.
        
        Args:
            source: File path or content.
            doc_type: Document type (pdf, docx, markdown). Auto-detected if None.
            **kwargs: Additional loader arguments.
            
        Returns:
            List of Document objects.
        """
        # Auto-detect type from extension if not provided
        if doc_type is None:
            path = Path(source) if isinstance(source, (str, Path)) else None
            if path and path.suffix:
                doc_type = path.suffix.lower().lstrip(".")
        
        if doc_type is None:
            raise ValueError("Document type could not be determined. Specify doc_type.")
        
        doc_type = doc_type.lower()
        
        if doc_type not in self._loaders:
            raise ValueError(
                f"Unsupported document type: {doc_type}. "
                f"Supported: {list(self._loaders.keys())}"
            )
        
        loader = self._loaders[doc_type]
        return loader.load(source, **kwargs)
    
    def load_markdown_content(
        self,
        content: str,
        source_name: str = "content",
    ) -> list[Document]:
        """
        Load markdown from raw content string.
        
        Args:
            content: Raw markdown content.
            source_name: Name to use as source.
            
        Returns:
            List of Document objects.
        """
        loader = self._loaders["markdown"]
        return loader.load(
            source=content,
            is_file=False,
            source_name=source_name,
        )
    
    @property
    def supported_types(self) -> list[str]:
        """Get list of supported document types."""
        return list(self._loaders.keys())
